"""
Enhanced document processor using PyMuPDF and pdfplumber.

This module provides layout-aware text extraction with bounding boxes,
table detection, and document structure preservation.
"""

import logging
from pathlib import Path
from typing import List, Dict, Any, Tuple
import fitz  # PyMuPDF
import pdfplumber
from datetime import datetime

logger = logging.getLogger(__name__)


class BoundingBox:
    """Represents a bounding box with coordinates."""

    def __init__(self, x1: float, y1: float, x2: float, y2: float):
        self.x1 = x1
        self.y1 = y1
        self.x2 = x2
        self.y2 = y2

    def to_list(self) -> List[float]:
        """Convert to list format [x1, y1, x2, y2]."""
        return [self.x1, self.y1, self.x2, self.y2]

    def to_dict(self) -> Dict[str, float]:
        """Convert to dictionary format."""
        return {"x1": self.x1, "y1": self.y1, "x2": self.x2, "y2": self.y2}

    @classmethod
    def from_rect(cls, rect) -> "BoundingBox":
        """Create from PyMuPDF Rect object."""
        return cls(rect.x0, rect.y0, rect.x1, rect.y1)

    @classmethod
    def merge(cls, boxes: List["BoundingBox"]) -> "BoundingBox":
        """Merge multiple bounding boxes into one encompassing box."""
        if not boxes:
            return cls(0, 0, 0, 0)

        x1 = min(box.x1 for box in boxes)
        y1 = min(box.y1 for box in boxes)
        x2 = max(box.x2 for box in boxes)
        y2 = max(box.y2 for box in boxes)

        return cls(x1, y1, x2, y2)


class TextBlock:
    """Represents a text block extracted from a document."""

    def __init__(
        self,
        text: str,
        bbox: BoundingBox,
        page_number: int,
        block_type: str = "text",
        confidence: float = 1.0,
        metadata: Dict[str, Any] = None
    ):
        self.text = text
        self.bbox = bbox
        self.page_number = page_number
        self.block_type = block_type  # text, table, header, footer, image
        self.confidence = confidence
        self.metadata = metadata or {}

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary representation."""
        return {
            "text": self.text,
            "bbox": self.bbox.to_list(),
            "page_number": self.page_number,
            "block_type": self.block_type,
            "confidence": self.confidence,
            "metadata": self.metadata
        }


class DocumentProcessor:
    """
    Enhanced document processor with layout awareness.

    Features:
    - Extract text blocks with bounding boxes (PyMuPDF)
    - Detect and extract tables (pdfplumber)
    - Preserve reading order
    - Identify block types (text, table, header, footer)
    """

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def process_pdf(self, file_path: Path) -> Tuple[List[TextBlock], Dict[str, Any]]:
        """
        Process a PDF file and extract text blocks with metadata.

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (text_blocks, document_metadata)
        """
        self.logger.info(f"Processing PDF: {file_path}")

        # Extract metadata
        metadata = self._extract_metadata(file_path)

        # Extract text blocks with PyMuPDF
        text_blocks = self._extract_text_blocks_pymupdf(file_path)

        # Extract tables with pdfplumber
        table_blocks = self._extract_tables_pdfplumber(file_path)

        # Merge text and table blocks, sorted by reading order
        all_blocks = self._merge_and_sort_blocks(text_blocks, table_blocks)

        self.logger.info(
            f"Extracted {len(all_blocks)} blocks "
            f"({len(text_blocks)} text, {len(table_blocks)} tables) "
            f"from {metadata['page_count']} pages"
        )

        return all_blocks, metadata

    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract document-level metadata."""
        try:
            doc = fitz.open(file_path)
            metadata = {
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "page_count": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "created": doc.metadata.get("creationDate", ""),
                "processed_at": datetime.now().isoformat()
            }
            doc.close()
            return metadata
        except Exception as e:
            self.logger.error(f"Error extracting metadata: {e}")
            return {
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "page_count": 0,
                "processed_at": datetime.now().isoformat()
            }

    def _extract_text_blocks_pymupdf(self, file_path: Path) -> List[TextBlock]:
        """
        Extract text blocks with bounding boxes using PyMuPDF.

        PyMuPDF's get_text("blocks") returns blocks in reading order with:
        - Block text
        - Bounding box coordinates
        - Block number (reading order)
        """
        blocks = []

        try:
            doc = fitz.open(file_path)

            for page_num in range(len(doc)):
                page = doc[page_num]

                # Get text blocks (format: x0, y0, x1, y1, "text", block_no, block_type)
                text_blocks = page.get_text("blocks")

                for block in text_blocks:
                    # block format: (x0, y0, x1, y1, "lines of text", block_no, block_type)
                    x0, y0, x1, y1, text, block_no, block_type = block

                    # Skip empty blocks
                    if not text.strip():
                        continue

                    # Create bounding box
                    bbox = BoundingBox(x0, y0, x1, y1)

                    # Determine block type (0 = text, 1 = image)
                    b_type = "image" if block_type == 1 else "text"

                    # Detect headers/footers (heuristic: top 10% or bottom 10% of page)
                    page_height = page.rect.height
                    if y0 < page_height * 0.1:
                        b_type = "header"
                    elif y1 > page_height * 0.9:
                        b_type = "footer"

                    text_block = TextBlock(
                        text=text.strip(),
                        bbox=bbox,
                        page_number=page_num + 1,  # 1-indexed
                        block_type=b_type,
                        confidence=1.0,
                        metadata={"block_number": block_no}
                    )

                    blocks.append(text_block)

            doc.close()

        except Exception as e:
            self.logger.error(f"Error extracting text blocks: {e}", exc_info=True)

        return blocks

    def _extract_tables_pdfplumber(self, file_path: Path) -> List[TextBlock]:
        """
        Extract tables using pdfplumber.

        Tables are converted to Markdown format and stored as table blocks.
        """
        table_blocks = []

        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Find tables on page
                    tables = page.find_tables()

                    for table_idx, table in enumerate(tables):
                        # Extract table data
                        table_data = table.extract()

                        if not table_data:
                            continue

                        # Convert table to Markdown format
                        markdown_table = self._table_to_markdown(table_data)

                        # Get table bounding box
                        bbox = BoundingBox(
                            table.bbox[0],  # x0
                            table.bbox[1],  # y0
                            table.bbox[2],  # x1
                            table.bbox[3]   # y1
                        )

                        table_block = TextBlock(
                            text=markdown_table,
                            bbox=bbox,
                            page_number=page_num + 1,
                            block_type="table",
                            confidence=0.95,
                            metadata={
                                "table_index": table_idx,
                                "rows": len(table_data),
                                "cols": len(table_data[0]) if table_data else 0,
                                "raw_data": table_data
                            }
                        )

                        table_blocks.append(table_block)

        except Exception as e:
            self.logger.error(f"Error extracting tables: {e}", exc_info=True)

        return table_blocks

    def _table_to_markdown(self, table_data: List[List[str]]) -> str:
        """Convert table data to Markdown format."""
        if not table_data:
            return ""

        lines = []

        # Header row
        header = table_data[0]
        lines.append("| " + " | ".join(str(cell or "") for cell in header) + " |")

        # Separator
        lines.append("| " + " | ".join("---" for _ in header) + " |")

        # Data rows
        for row in table_data[1:]:
            lines.append("| " + " | ".join(str(cell or "") for cell in row) + " |")

        return "\n".join(lines)

    def _merge_and_sort_blocks(
        self,
        text_blocks: List[TextBlock],
        table_blocks: List[TextBlock]
    ) -> List[TextBlock]:
        """
        Merge text and table blocks, sorted by page and position.

        Sorting: by page number, then by y-coordinate (top to bottom),
        then by x-coordinate (left to right).
        """
        all_blocks = text_blocks + table_blocks

        # Sort by page, then y-coordinate, then x-coordinate
        all_blocks.sort(key=lambda b: (b.page_number, b.bbox.y1, b.bbox.x1))

        return all_blocks

    def process_docx(self, file_path: Path) -> Tuple[List[TextBlock], Dict[str, Any]]:
        """
        Process a DOCX file.

        Note: DOCX doesn't have page-based layout, so bounding boxes are simulated.
        """
        from docx import Document

        blocks = []

        try:
            doc = Document(file_path)

            # Simulate page number (rough estimate: 500 words per page)
            words_per_page = 500
            word_count = 0
            page_num = 1

            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()

                if not text:
                    continue

                # Update page number based on word count
                word_count += len(text.split())
                page_num = (word_count // words_per_page) + 1

                # Simulated bounding box (no real coordinates in DOCX)
                bbox = BoundingBox(0, para_idx * 20, 612, (para_idx + 1) * 20)

                block = TextBlock(
                    text=text,
                    bbox=bbox,
                    page_number=page_num,
                    block_type="text",
                    confidence=1.0,
                    metadata={"paragraph_index": para_idx}
                )

                blocks.append(block)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                markdown_table = self._table_to_markdown(table_data)

                # Update page number
                word_count += len(markdown_table.split())
                page_num = (word_count // words_per_page) + 1

                bbox = BoundingBox(0, 0, 612, 792)  # Default page size

                table_block = TextBlock(
                    text=markdown_table,
                    bbox=bbox,
                    page_number=page_num,
                    block_type="table",
                    confidence=0.9,
                    metadata={
                        "table_index": table_idx,
                        "rows": len(table_data),
                        "cols": len(table_data[0]) if table_data else 0
                    }
                )

                blocks.append(table_block)

            metadata = {
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "page_count": page_num,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "processed_at": datetime.now().isoformat()
            }

        except Exception as e:
            self.logger.error(f"Error processing DOCX: {e}", exc_info=True)
            blocks = []
            metadata = {
                "file_name": file_path.name,
                "error": str(e)
            }

        return blocks, metadata
